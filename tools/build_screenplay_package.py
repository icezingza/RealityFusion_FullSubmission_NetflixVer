#!/usr/bin/env python3
"""Build current Reality Fusion screenplay drafts into DOCX/PDF packages.

This script reads only the active draft set declared below. It does not modify
source screenplays and does not claim production lock or validated runtime.
"""

from __future__ import annotations

import re
import shutil
import subprocess
import sys
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "dist" / "screenplay-package"
DOCX_DIR = OUT / "docx"
PDF_DIR = OUT / "pdf"

ACTIVE_DRAFTS = [
    ("EP1", ROOT / "screenplays/EP1/Reality_Fusion_EP1_Master_Draft_v2.1.md"),
    ("EP2", ROOT / "screenplays/EP2/Reality_Fusion_EP2_Master_Draft_v2.1.md"),
    ("EP3", ROOT / "screenplays/EP3/Reality_Fusion_EP3_Master_Draft_v2.2.md"),
    ("EP4", ROOT / "screenplays/EP4/Reality_Fusion_EP4_Master_Draft_v2.1.md"),
    ("EP5", ROOT / "screenplays/EP5/Reality_Fusion_EP5_Master_Draft_v2.1.md"),
]

SCENE_RE = re.compile(r"^###\s+(\d+)\.\s+(.+)$")
CHAR_RE = re.compile(r"^[\wก-๙ .()\-/]+$")
TRANSITIONS = {"CUT TO:", "FADE IN:", "FADE OUT.", "CUT TO BLACK."}


def set_run_font(run, name: str, size: float, bold: bool = False, italic: bool = False) -> None:
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), name)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)
    set_run_font(run, "Noto Sans Thai", 9)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.2)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.8)
    add_page_number(section.footer.paragraphs[0])

    normal = doc.styles["Normal"]
    normal.font.name = "Noto Sans Thai"
    normal.font.size = Pt(11.5)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.0


def add_para(doc: Document, text: str = "", *, align=None, left=0, right=0,
             before=0, after=0, keep=False, bold=False, italic=False,
             font="Noto Sans Thai", size=11.5) -> None:
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.left_indent = Cm(left)
    pf.right_indent = Cm(right)
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.keep_with_next = keep
    run = p.add_run(text)
    set_run_font(run, font, size, bold=bold, italic=italic)


def is_character_cue(line: str, next_line: str | None) -> bool:
    if not line or line.startswith(("#", "`", "- ")):
        return False
    if line in TRANSITIONS:
        return False
    if len(line) > 45 or not CHAR_RE.match(line):
        return False
    if next_line is None or not next_line.strip():
        return False
    markers = ("INT.", "EXT.", "LESSON", "TERM:", "STATUS:", "OBSERVATIONS:", "UNRESOLVED:")
    if line.startswith(markers):
        return False
    return True


def build_docx(source: Path, destination: Path) -> None:
    text = source.read_text(encoding="utf-8")
    lines = text.splitlines()
    doc = Document()
    configure_document(doc)

    in_code = False
    code_lines: list[str] = []
    i = 0
    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()
        stripped = line.strip()
        next_line = lines[i + 1].strip() if i + 1 < len(lines) else None

        if stripped.startswith("```"):
            if not in_code:
                in_code = True
                code_lines = []
            else:
                in_code = False
                for code_line in code_lines:
                    add_para(doc, code_line, left=1.0, right=0.6, font="Noto Sans Mono", size=9.5)
                add_para(doc, "", after=2)
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if not stripped or stripped == "---":
            add_para(doc, "", after=1)
            i += 1
            continue

        if stripped.startswith("# "):
            add_para(doc, stripped[2:], align=WD_ALIGN_PARAGRAPH.CENTER,
                     before=42, after=8, bold=True, size=18)
        elif stripped.startswith("## "):
            add_para(doc, stripped[3:], align=WD_ALIGN_PARAGRAPH.CENTER,
                     after=10, bold=True, size=14)
        elif stripped.startswith("**") and stripped.endswith("**"):
            add_para(doc, stripped.strip("* "), align=WD_ALIGN_PARAGRAPH.CENTER,
                     after=4, bold=True, size=10.5)
        elif (m := SCENE_RE.match(stripped)):
            add_para(doc, f"{m.group(1)}. {m.group(2)}", before=10, after=5,
                     keep=True, bold=True, size=11.5)
        elif stripped in TRANSITIONS:
            add_para(doc, stripped, align=WD_ALIGN_PARAGRAPH.RIGHT,
                     before=6, after=6, bold=True, size=10.5)
        elif stripped.startswith("`") and stripped.endswith("`"):
            add_para(doc, stripped.strip("`"), left=1.0, right=0.6,
                     font="Noto Sans Mono", size=9.5)
        elif is_character_cue(stripped, next_line):
            add_para(doc, stripped, left=5.1, right=1.7, before=5, keep=True,
                     bold=True, size=10.5)
        elif i > 0 and is_character_cue(lines[i - 1].strip(), stripped):
            add_para(doc, stripped, left=3.4, right=3.0, after=2, size=11)
        else:
            clean = stripped.replace("**", "")
            add_para(doc, clean, after=2, size=11.5)
        i += 1

    destination.parent.mkdir(parents=True, exist_ok=True)
    doc.save(destination)


def convert_to_pdf(docx_path: Path, pdf_dir: Path) -> Path:
    pdf_dir.mkdir(parents=True, exist_ok=True)
    cmd = [
        "libreoffice", "--headless", "--convert-to", "pdf",
        "--outdir", str(pdf_dir), str(docx_path),
    ]
    subprocess.run(cmd, check=True)
    pdf = pdf_dir / f"{docx_path.stem}.pdf"
    if not pdf.exists() or pdf.stat().st_size == 0:
        raise RuntimeError(f"PDF conversion failed: {docx_path}")
    return pdf


def main() -> int:
    missing = [str(path) for _, path in ACTIVE_DRAFTS if not path.exists()]
    if missing:
        print("Missing active drafts:", *missing, sep="\n- ", file=sys.stderr)
        return 2

    if OUT.exists():
        shutil.rmtree(OUT)
    DOCX_DIR.mkdir(parents=True)
    PDF_DIR.mkdir(parents=True)

    built: list[Path] = []
    for episode, source in ACTIVE_DRAFTS:
        docx_name = source.with_suffix(".docx").name
        docx_path = DOCX_DIR / docx_name
        build_docx(source, docx_path)
        pdf_path = convert_to_pdf(docx_path, PDF_DIR)
        built.extend([docx_path, pdf_path])
        print(f"Built {episode}: {docx_path.name}, {pdf_path.name}")

    readme = OUT / "README.txt"
    readme.write_text(
        "REALITY FUSION - CURRENT SCREENPLAY FORMATTING PACKAGE\n\n"
        "Active drafts: EP1 v2.1, EP2 v2.1, EP3 v2.2, EP4 v2.1, EP5 v2.1.\n"
        "These files are formatting outputs for table read, page review, and specialist review.\n"
        "They are not evidence of production lock, validated runtime, or external specialist approval.\n",
        encoding="utf-8",
    )
    built.append(readme)

    zip_path = OUT / "Reality_Fusion_Current_Screenplays_DOCX_PDF.zip"
    with ZipFile(zip_path, "w", ZIP_DEFLATED) as archive:
        for path in built:
            archive.write(path, path.relative_to(OUT))
    print(f"Package: {zip_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
